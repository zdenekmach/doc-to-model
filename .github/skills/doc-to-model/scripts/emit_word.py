#!/usr/bin/env python3
"""emit_word.py — deterministický emitor: instance analytical-doc → .docx

Čte jednu instanci strukturované pravdy a vygeneruje strukturovaný Word dokument.
Stejný vstup = stejný výstup, žádné generování textu.

Každá sekce je volitelná — emitor vypíše jen to, co v modelu je.

    python3 emit_word.py --model model.yaml --out out/dokument.docx
"""
import argparse
import sys
from pathlib import Path

import yaml

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import lang
from pipeline_state import require, warn_if_missing

BRAND = RGBColor(0x1F, 0x49, 0x7D)
PRIORITY_ORDER = {"MUST": 0, "SHOULD": 1, "COULD": 2, "WONT": 3}

# Popisky Wordu. Word je PROJEKCE modelu, takže má mluvit jazykem zdroje —
# česká hlavička nad anglickým obsahem dokument znehodnocuje pro toho, komu se
# posílá. Prohlížeč je naopak nástroj, ne výstup, a zůstává česky.
_PACK = None
_REF = None


def L(key: str) -> str:
    """Popisek v jazyce modelu."""
    if _PACK is None:
        raise RuntimeError("jazykový balíček nebyl nastaven — viz main()")
    return lang.label(_PACK, key, _REF)


def confidence_label(value: str) -> str:
    """Stav jistoty výroku v jazyce modelu."""
    return {
        "explicit": L("conf_explicit"),
        "derived": L("conf_derived"),
        "assumed": L("assumed"),
    }.get(value, value)


def load(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


def index_by_id(items):
    return {it["id"]: it for it in (items or []) if "id" in it}


def kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        if v in (None, "", []):
            continue
        cells = t.add_row().cells
        cells[0].text = str(k)
        cells[1].text = "\n".join(v) if isinstance(v, list) else str(v)
        for p in cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
    return t


def header_table(doc, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    return t


def source_label(src_id, sources):
    src = sources.get(src_id)
    if not src:
        return src_id or ""
    parts = [src.get("title", src_id)]
    if src.get("locator"):
        parts.append(src["locator"])
    return ", ".join(parts)


def build(m: dict) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    sources = index_by_id(m.get("sources"))
    actors = index_by_id(m.get("actors"))

    title = doc.add_heading(m.get("title", L("untitled")), level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND

    if m.get("subtitle"):
        sub = doc.add_paragraph(m["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if sub.runs:
            sub.runs[0].italic = True

    kv_table(doc, [
        ("ID", m.get("id")),
        (L("version"), m.get("version")),
        (L("status"), m.get("status")),
        (L("author"), m.get("author")),
        (L("date"), m.get("date")),
        (L("source_document"), m.get("source_document")),
        (L("regulations"), m.get("regulations")),
    ])
    doc.add_paragraph()

    n = 0

    if m.get("context"):
        n += 1
        doc.add_heading(f"{n}. {L('context')}", level=1)
        doc.add_paragraph(str(m["context"]).strip())

    if m.get("actors"):
        n += 1
        doc.add_heading(f"{n}. {L('actors')}", level=1)
        t = header_table(doc, ["ID", L("actor"), L("role")])
        for a in m["actors"]:
            c = t.add_row().cells
            c[0].text = a.get("id", "")
            c[1].text = a.get("name", "")
            c[2].text = a.get("role", "")

    if m.get("scope_in") or m.get("scope_out"):
        n += 1
        doc.add_heading(f"{n}. {L('scope')}", level=1)
        if m.get("scope_in"):
            doc.add_heading(L("scope_in"), level=2)
            for item in m["scope_in"]:
                doc.add_paragraph(item, style="List Bullet")
        if m.get("scope_out"):
            doc.add_heading(L("scope_out"), level=2)
            for item in m["scope_out"]:
                doc.add_paragraph(item, style="List Bullet")

    if m.get("claims"):
        n += 1
        doc.add_heading(f"{n}. {L('claims')}", level=1)
        for c in m["claims"]:
            head = f"{c.get('id', '')} — {c.get('title', '')}"
            if c.get("claim_type"):
                head += f"  [{c['claim_type']}]"
            doc.add_heading(head, level=2)
            if c.get("description"):
                doc.add_paragraph(str(c["description"]).strip())
            if c.get("basis"):
                p = doc.add_paragraph(f"{L('basis')}: {c['basis']}")
                if p.runs:
                    p.runs[0].italic = True
            meta_bits = []
            if c.get("scope"):
                meta_bits.append(f"{L('scope_of_validity')}: {c['scope']}")
            if c.get("source"):
                meta_bits.append(f"{L('source')}: {source_label(c['source'], sources)}")
            if c.get("confidence"):
                meta_bits.append(confidence_label(c["confidence"]))
            if meta_bits:
                p = doc.add_paragraph(" · ".join(meta_bits))
                if p.runs:
                    p.runs[0].italic = True

    if m.get("requirements"):
        n += 1
        doc.add_heading(f"{n}. {L('requirements')}", level=1)
        reqs = sorted(
            m["requirements"],
            key=lambda r: (PRIORITY_ORDER.get(r.get("priority"), 9), r.get("id", "")),
        )
        for r in reqs:
            head = f"{r.get('id', '')} — {r.get('title', '')}"
            if r.get("priority"):
                head += f"  [{r['priority']}]"
            doc.add_heading(head, level=2)
            if r.get("description"):
                doc.add_paragraph(str(r["description"]).strip())
            if r.get("acceptance"):
                p = doc.add_paragraph(L("acceptance") + ":")
                p.runs[0].bold = True
                for ac in r["acceptance"]:
                    doc.add_paragraph(ac, style="List Bullet")
            if r.get("justified_by"):
                claims_by_id = index_by_id(m.get("claims"))
                duvody = "; ".join(
                    f"{j} — {claims_by_id.get(j, {}).get('title', '?')}"
                    for j in r["justified_by"]
                )
                p = doc.add_paragraph(f"{L('justified_by')}: {duvody}")
                if p.runs:
                    p.runs[0].italic = True
            meta_bits = []
            if r.get("actor"):
                who = actors.get(r["actor"], {}).get("name", r["actor"])
                meta_bits.append(f"{L('responsible')}: {who}")
            if r.get("source"):
                meta_bits.append(f"{L('source')}: {source_label(r['source'], sources)}")
            if r.get("confidence"):
                meta_bits.append(confidence_label(r["confidence"]))
            if meta_bits:
                p = doc.add_paragraph(" · ".join(meta_bits))
                if p.runs:
                    p.runs[0].italic = True

    if m.get("quality_requirements"):
        n += 1
        doc.add_heading(f"{n}. {L('quality_requirements')}", level=1)
        t = header_table(doc, ["ID", L("category"), L("requirement"), L("source")])
        for q in m["quality_requirements"]:
            c = t.add_row().cells
            c[0].text = q.get("id", "")
            c[1].text = q.get("category", "")
            c[2].text = str(q.get("requirement", "")).strip()
            c[3].text = source_label(q.get("source"), sources)

    if m.get("entities"):
        n += 1
        doc.add_heading(f"{n}. {L('entities')}", level=1)
        for e in m["entities"]:
            doc.add_heading(f"{e.get('name', e.get('id', ''))}", level=2)
            if e.get("description"):
                doc.add_paragraph(e["description"])
            if e.get("fields"):
                t = header_table(doc, [L("field"), L("type"), L("required"), L("note")])
                for f in e["fields"]:
                    c = t.add_row().cells
                    c[0].text = f.get("name", "")
                    c[1].text = f.get("type", "")
                    c[2].text = "Ano" if f.get("required_field") else "Ne"
                    c[3].text = f.get("note", "")

    if m.get("process"):
        n += 1
        doc.add_heading(f"{n}. {L('process')}", level=1)
        proc = m["process"]
        doc.add_paragraph(proc.get("name", ""))
        t = header_table(doc, [L("step"), L("type"), L("description"), L("actor"), L("next")])
        for s in proc.get("steps", []):
            c = t.add_row().cells
            c[0].text = s.get("id", "")
            c[1].text = s.get("type", "")
            c[2].text = s.get("label", "")
            c[3].text = actors.get(s.get("actor"), {}).get("name", s.get("actor", "") or "")
            if s.get("branches"):
                c[4].text = "; ".join(
                    f"{b.get('label')} → {b.get('target')}" for b in s["branches"]
                )
            else:
                c[4].text = s.get("next", "") or ""
        p = doc.add_paragraph(
            L("diagram_note")
        )
        if p.runs:
            p.runs[0].italic = True

    if m.get("risks"):
        n += 1
        doc.add_heading(f"{n}. {L('risks')}", level=1)
        t = header_table(doc, ["ID", L("risk"), L("mitigation")])
        for r in m["risks"]:
            c = t.add_row().cells
            c[0].text = r.get("id", "")
            c[1].text = r.get("description", "")
            c[2].text = r.get("mitigation", "") or ""

    if m.get("open_questions"):
        n += 1
        doc.add_heading(f"{n}. {L('open_questions')}", level=1)
        for q in m["open_questions"]:
            doc.add_paragraph(q, style="List Bullet")

    if sources:
        n += 1
        doc.add_heading(f"{n}. {L('sources')}", level=1)
        t = header_table(doc, ["ID", L("source"), L("locator")])
        for sid, s in sources.items():
            c = t.add_row().cells
            c[0].text = sid
            c[1].text = s.get("title", "")
            c[2].text = s.get("locator", "") or s.get("url", "") or ""

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = (
        f"{m.get('id', '')} v{m.get('version', '')} · " + L('footer')
    )
    return doc


def main():
    ap = argparse.ArgumentParser(description="Emit .docx z instance analytical-doc.")
    ap.add_argument("--model", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    ap.add_argument("--lang", help="Kód jazyka popisků (default: detekce z modelu).")
    args = ap.parse_args()

    require(args.model, ["validate"], "emit")
    warn_if_missing(args.model, "ground-check",
                    "výstup může nést tvrzení bez opory ve zdroji")

    m = load(args.model)

    # Jazyk se bere z MODELU, ne ze zdrojového textu — emitor zdroj nedostává
    # a model je psaný jazykem dokumentu, takže nese stejnou informaci.
    global _PACK, _REF
    packs = lang.load_packs()
    _REF = packs["cs"].labels if "cs" in packs else {}
    model_text = " ".join(
        str(x.get(k, ""))
        for coll in ("claims", "requirements", "quality_requirements")
        for x in (m.get(coll) or [])
        for k in ("title", "description")
    ) or str(m.get("title", ""))
    _PACK, how = lang.resolve(model_text, args.lang)
    print(f"[jazyk] popisky: {how}")

    doc = build(m)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)

    # Počty, ne jen cesta. Word z jednoho požadavku vypadal v logu stejně jako
    # ze dvou set — a tenhle řádek je často jediné, co člověk z běhu přečte.
    reqs = len(m.get("requirements") or []) + len(m.get("quality_requirements") or [])
    claims = len(m.get("claims") or [])
    steps = len((m.get("process") or {}).get("steps") or [])
    print(f"[OK] Word: {args.out} "
          f"({reqs} požadavků · {claims} tvrzení · {steps} kroků procesu)")

    if not reqs and not claims:
        print("  [POZOR] dokument nenese žádný požadavek ani tvrzení — "
              "vznikl formálně správný, ale prázdný výstup.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
