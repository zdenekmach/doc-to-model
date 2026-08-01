#!/usr/bin/env python3
"""ingest.py — dokument → text, se značkami, na které se dá odkazovat.

Vstupem je to, co reálně chodí: PDF, Word, markdown, holý text. Výstupem jeden
`.txt`, ve kterém jsou zachovány záchytné body — u PDF hranice stran, u Wordu
nadpisy. Bez nich by se pak nedalo citovat na místo, jen na dokument.

    python3 ingest.py --in dokument.pdf --out zdroj.txt

Skenované PDF bez textové vrstvy pozná podle mizivého výtěžku a řekne to nahlas
— tichý prázdný výstup je horší než chyba.
"""
import argparse
import sys
from pathlib import Path

PAGE_MARK = "=== STRANA {n} ==="
# Pod tuhle hranici znaků na stránku to vypadá na sken bez textové vrstvy.
THIN_PAGE = 80


def from_pdf(path: Path):
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.exit("[CHYBA] Chybí pypdf. Nainstaluj: pip install pypdf")
    reader = PdfReader(str(path))
    parts, thin = [], 0
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if len(text) < THIN_PAGE:
            thin += 1
        parts.append(f"\n{PAGE_MARK.format(n=i)}\n{text}")
    note = None
    if reader.pages and thin / len(reader.pages) > 0.5:
        note = (f"{thin} z {len(reader.pages)} stran nemá skoro žádný text — "
                "nejspíš sken bez textové vrstvy, potřebuje OCR")
    return "\n".join(parts), len(reader.pages), note


def from_docx(path: Path):
    try:
        from docx import Document
    except ImportError:
        sys.exit("[CHYBA] Chybí python-docx. Nainstaluj: pip install python-docx")
    doc = Document(str(path))
    out = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            out.append("")
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            # Úroveň nadpisu → markdown, aby ji segment.py uměl najít.
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            out.append(f"\n{'#' * min(int(level), 6)} {text}")
        else:
            out.append(text)
    for t_i, table in enumerate(doc.tables, 1):
        out.append(f"\n[tabulka {t_i}]")
        for row in table.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out), len(doc.paragraphs), None


def from_text(path: Path):
    text = path.read_text(encoding="utf-8")
    return text, text.count("\n") + 1, None


READERS = {
    ".pdf": from_pdf,
    ".docx": from_docx,
    ".md": from_text,
    ".txt": from_text,
    ".markdown": from_text,
}


def main():
    ap = argparse.ArgumentParser(description="Převeď dokument na text se záchytnými body.")
    ap.add_argument("--in", dest="src", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    args = ap.parse_args()

    if not args.src.exists():
        sys.exit(f"[CHYBA] Soubor nenalezen: {args.src}")
    reader = READERS.get(args.src.suffix.lower())
    if not reader:
        sys.exit(f"[CHYBA] Neznámý formát: {args.src.suffix}. "
                 f"Umím: {', '.join(sorted(READERS))}")

    text, units, note = reader(args.src)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(text, encoding="utf-8")

    print(f"[ingest] {args.src.name} → {args.out} "
          f"({len(text)} znaků, {units} stran/odstavců)")
    if note:
        print(f"[POZOR] {note}")
    if len(text.strip()) < 200:
        sys.exit("[CHYBA] Z dokumentu skoro nic nevypadlo — extrakce se nepovedla.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
